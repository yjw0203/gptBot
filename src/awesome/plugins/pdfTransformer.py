import uuid
import PyPDF2
import docx
from nonebot import on_notice, NoticeSession
from aiocqhttp import Message
import requests
import openai
import os

@on_notice('offline_file')
async def pdf(session: NoticeSession):
    file = session.event['file']
    if file['name'].endswith(".pdf"):
        try:
            await session.send('正在下载文件：{}！'.format(file['name']))
            new_fileName_pdf = await downloadFile(file['url'],file['name'])
            await session.send('文件下载成功：{}！'.format(new_fileName_pdf))
            await session.send('正在翻译文件：{}！'.format(new_fileName_pdf))
            new_fileName_docx = await pdfTransform(new_fileName_pdf,session)
            await session.send('翻译完成：{}！'.format(new_fileName_docx))
            await session.send('现在将文件发送给你：{}！'.format(new_fileName_docx))
            await session.bot.call_action('upload_private_file',user_id=session.event.user_id,file=os.getcwd() +'\\' + new_fileName_docx,name=new_fileName_docx)
            await session.send('pdf{}翻译完毕！'.format(file['name']))
            os.remove(os.getcwd() +'\\' + new_fileName_pdf)
            os.remove(os.getcwd() +'\\' + new_fileName_docx)
        except Exception as err:
            await session.send('翻译失败:{}'.format(file['name']))
    
async def downloadFile(url,fileName):
    new_fileName = str(uuid.uuid4().hex) + fileName
    r = requests.get(url)
    with open(new_fileName,"wb") as f:
        f.write(r.content)
    return new_fileName

async def pdfTransform(fileName,session):
    pdf_file = open(fileName, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    doc = docx.Document()
    page_index = 1
    for page in pdf_reader.pages:
        text = page.extract_text()
        transform_text = await transform(text)
        paragraph = doc.add_paragraph(transform_text)
        paragraph.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        await session.send('第{}页翻译完成：{}！'.format(page_index,fileName))
        page_index += 1
    new_fileName = fileName[:-4] + ".docx"
    doc.save(new_fileName)
    pdf_file.close()
    return new_fileName

async def transform(text):
    try:
        text = '\'' + text + '\' 翻译成中文,不要复述原话'
        completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo", 
                messages=[{"role": "system", "content": text}],
                timeout = 5000
        )
        response = completion.choices[0].message.content.strip()
    except Exception as err:
        return '$翻译异常$'
    return response